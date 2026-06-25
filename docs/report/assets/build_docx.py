#!/usr/bin/env python3
"""Generate the MEDIS OOAD report as a professionally formatted .docx."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = "/Users/huichiyap/Documents/OOAD MEDIS/MEDIS/Medical-Information-Scheduling-MEDIS-"
ASSETS = os.path.join(ROOT, "docs", "report", "assets")
OUT = os.path.join(ROOT, "docs", "report", "MEDIS-Report.docx")

BLUE = RGBColor(0x1F, 0x3A, 0x8A)
DARK = RGBColor(0x1E, 0x29, 0x3B)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, sz in [(1, 15), (2, 13), (3, 12)]:
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = BLUE
    st.font.bold = True

# page margins
for s in doc.sections:
    s.top_margin = Inches(1); s.bottom_margin = Inches(1)
    s.left_margin = Inches(1); s.right_margin = Inches(1)


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_field(paragraph, instruction):
    """Insert a Word field (e.g. PAGE, NUMPAGES, TOC)."""
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = ""
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(b); r.append(instr); r.append(sep); r.append(t); r.append(e)


def para(text="", *, bold=False, italic=False, size=None, color=None,
         align=None, style=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = color
    if align is not None: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def code_block(code):
    for line in code.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F4F7")
        pPr.append(shd)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.font.size = Pt(10)
        set_cell_bg(hdr[i], "1F3A8A")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def figure(img, caption, width=6.3):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(ASSETS, img), width=Inches(width))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY


def h1(text): doc.add_heading(text, level=1)
def h2(text): doc.add_heading(text, level=2)
def pagebreak(): doc.add_page_break()


# ======================= COVER PAGE =======================
doc.add_paragraph().paragraph_format.space_after = Pt(60)
para("MEDIS", bold=True, size=40, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Hospital Management System", bold=True, size=22, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Object-Oriented Analysis and Design — Project Report", size=13, italic=True,
     color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
para("CCP6224 — Object-Oriented Analysis and Design", bold=True, size=13,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Lab Exercise (Group Project)", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Technology: Java Swing + SQLite (JDBC)", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

table(["Field", "Detail"],
      [["Tutorial Section", "[Tutorial Section]"],
       ["Lecturer / Tutor", "[Lecturer / Tutor Name]"],
       ["Submission Date", "[Submission Date]  (Due: 29 June 2026)"]],
      widths=[2.3, 4.0])

para("Group Members", bold=True, size=12, space_after=4)
table(["No.", "Student Name", "Student ID", "Role / Contribution"],
      [["1", "[Student Name 1]", "[Student ID 1]", "[Contribution]"],
       ["2", "[Student Name 2]", "[Student ID 2]", "[Contribution]"],
       ["3", "[Student Name 3]", "[Student ID 3]", "[Contribution]"],
       ["4", "[Student Name 4]", "[Student ID 4]", "[Contribution]"]],
      widths=[0.6, 2.4, 1.6, 1.9])
para("Declaration: This is original group work produced solely by the members listed "
     "above, in accordance with the assignment integrity rules.", italic=True, size=10, color=GREY)
pagebreak()

# ======================= TABLE OF CONTENTS =======================
h1("Table of Contents")
toc_p = doc.add_paragraph()
add_field(toc_p, 'TOC \\o "1-2" \\h \\z \\u')
para('(If the contents do not appear, click here and press F9, or choose '
     '"Update Field" — Word builds the page numbers automatically.)',
     italic=True, size=9, color=GREY)
pagebreak()

# ======================= 1. INTRODUCTION =======================
h1("1. Introduction")
para("This report documents the analysis, design, and implementation of MEDIS, a desktop "
     "Hospital Management System (HMS) developed for the CCP6224 Object-Oriented Analysis "
     "and Design lab exercise. The system is built entirely with Java Swing for the "
     "graphical user interface and SQLite (accessed through JDBC) for persistent storage, "
     "in line with the assignment's instruction to use Java Swing only and avoid external "
     "application frameworks.")
para("The objective of the assignment is to extend prior single-feature Swing exercises "
     "into a cohesive, multi-module application that demonstrates strong object-oriented "
     "design, proper UML modelling, and a functional GUI. MEDIS satisfies these goals "
     "through five integrated modules — User Login, Appointment Booking, Patient Record "
     "Management, Doctor Management, and Basic Reporting — organised under a clean, "
     "layered architecture.")
para("This report explains the system's requirements, presents the supporting UML models "
     "(one Use Case diagram, one Class diagram, and four Sequence diagrams), maps each "
     "object-oriented principle directly to the implemented source code, and summarises "
     "the delivered features. All diagrams are derived from the actual submitted source "
     "code to guarantee design–code coherence, an explicit criterion in the marking rubric.")

# ======================= 2. PROJECT BACKGROUND =======================
h1("2. Project Background")
para("Hospitals coordinate large volumes of information — patient records, doctor "
     "specialisations, and appointment schedules — that must remain consistent and "
     "accessible to staff in different roles. A manual or paper-based process is "
     "error-prone: double-booked time slots, lost patient histories, and uncontrolled "
     "access to records are common problems.")
para("MEDIS addresses a simplified version of this domain. It provides three categories "
     "of staff — Admin, Doctor, and Receptionist — with role-appropriate access to the "
     "system. Receptionists register patients and book appointments, doctors review their "
     "schedules and patient histories, and administrators oversee all modules including "
     "reporting. The system enforces core business rules such as input validation, secure "
     "password storage, and prevention of duplicate appointment slots.")
para("The project deliberately models a real-world scenario rather than a toy example, in "
     "keeping with the assignment's emphasis on real-world OO modelling and the expansion "
     "from single-feature apps to a multi-module system.")

# ======================= 3. SYSTEM OVERVIEW =======================
h1("3. System Overview")
para("MEDIS is a single-user desktop application launched from the Main class. On first "
     "launch it auto-creates and seeds a local SQLite database (db/medis.db) from "
     "db/schema.sql. The user authenticates through a login window; upon success a "
     "role-based dashboard presents only the tabs the user is permitted to use.")
para("The system follows a strict four-layer architecture — a Model–View–Controller (MVC) "
     "arrangement augmented with a Data Access Object (DAO) layer for persistence. Control "
     "flows downward (View → Controller → DAO → Database) and results flow back upward, "
     "with plain model objects carrying data across every layer.")
code_block("view/  -->  controller/  -->  dao/  -->  db/  -->  SQLite (db/medis.db)\n"
           "(Swing GUI)  (validation +   (interface    (singleton\n"
           "             orchestration)   + JDBC impl)   connection)\n\n"
           "model/ : Patient, Doctor, Appointment, User, Result  (flow across all layers)\n"
           "util/  : Validator, PasswordHasher  (cross-cutting helpers)")
para("A central SystemController acts as the single wiring hub: it holds every "
     "sub-controller and the currently logged-in User, and is the only object handed to "
     "the Swing frames. This keeps the views free of any direct database or DAO "
     "instantiation.")
para("Design note. Unlike the early planning brief in docs/design/diagrams.md, the final "
     "implementation contains no separate \"Service\" layer. Business logic (validation, "
     "duplicate-slot checks, orchestration) lives in the controllers, and all SQL lives "
     "in the DAO implementations. The UML models in this report reflect the actual "
     "delivered code, ensuring full design–code consistency.", italic=True, size=10, color=GREY)

# ======================= 4. FUNCTIONAL REQUIREMENTS =======================
h1("4. Functional Requirements")
para("The assignment specifies five mandatory modules. The table below maps each "
     "requirement to its implementing classes in the submitted source.")
table(["Module", "Required Functions", "Implementation"],
      [["M1 — User Login", "Username/password login; role-based access; input validation",
        "LoginController.authenticate(), LoginFrame, DashboardFrame, Validator, PasswordHasher"],
       ["M2 — Appointment Booking", "Create appointment; select patient/doctor/date-time; view list; prevent duplicate slots",
        "AppointmentController.book()/cancel()/getAll(), AppointmentPanel, AppointmentDAOImpl"],
       ["M3 — Patient Records", "Add patient; update details; view history; store name/age/gender/history",
        "PatientController.add()/update()/getAll(), PatientPanel, PatientDAOImpl"],
       ["M4 — Doctor Management", "Add/view doctors; assign specialisation; link to appointments",
        "DoctorController.add()/getAll(), DoctorPanel, DoctorDAOImpl"],
       ["M5 — Basic Reporting", "Total patients; total appointments; doctor schedules",
        "ReportController.totalPatients()/totalAppointments()/doctorSchedules(), ReportPanel"]],
      widths=[1.5, 2.6, 2.4])
para("Non-functional / technical requirements satisfied:", bold=True, space_after=2)
bullet("Java Swing only — the entire GUI uses JFrame, JPanel, JTable, JComboBox, JOptionPane, etc.; no external UI framework.")
bullet("Event handling — every button is wired with an ActionListener lambda (e.g. book.addActionListener(e -> doBook())).")
bullet("Input validation — centralised in util.Validator (blank checks, age 0–150, future-date) and invoked by the controllers.")
bullet("Security — passwords are stored as SHA-256 hashes (util.PasswordHasher); the database seed stores hashes, never plaintext.")
bullet("Clean architecture — strict layer separation with DAO interfaces decoupling controllers from JDBC.")

# ======================= 5. OBJECT-ORIENTED DESIGN =======================
h1("5. Object-Oriented Design")
para("The design organises responsibilities into cohesive packages, each with a single concern:")
table(["Package", "Responsibility", "Key Classes"],
      [["model", "Plain data carriers (entities) and the Result envelope",
        "User (abstract), Admin, Doctor, Receptionist, Patient, Appointment, Result"],
       ["view", "Swing GUI and dialogs",
        "LoginFrame, DashboardFrame, PatientPanel, DoctorPanel, AppointmentPanel, ReportPanel, DialogHelper"],
       ["controller", "Validation, business rules, orchestration",
        "SystemController, LoginController, PatientController, DoctorController, AppointmentController, ReportController"],
       ["dao", "Data access (interface + JDBC implementation pairs)",
        "UserDAO/Impl, PatientDAO/Impl, DoctorDAO/Impl, AppointmentDAO/Impl"],
       ["db", "Database connection management", "DatabaseConnection (singleton)"],
       ["util", "Cross-cutting helpers", "Validator, PasswordHasher"]],
      widths=[1.1, 2.5, 2.9])
para("Key design decisions and their rationale:", bold=True, space_after=2)
para("1. Controllers never throw exceptions to the view. Every controller method returns a "
     "model.Result object (Result.ok(), Result.ok(data), or Result.fail(message)). Views "
     "check r.isOk() and display r.getMessage(). This produces a uniform, predictable "
     "error-handling contract across the whole system.")
para("2. DAOs are interface + implementation pairs. Controllers depend on the interface "
     "(e.g. PatientDAO), not the JDBC implementation. This is the abstraction seam that "
     "would allow the SQLite backend to be replaced without touching controller code.")
para("3. SystemController is the composition root. Main builds the object graph "
     "(DAOs → controllers → SystemController) once and passes the single SystemController "
     "to each frame, avoiding scattered object creation inside the GUI.")
para("4. A single canonical date-time format. AppointmentDAOImpl.FMT (yyyy-MM-dd HH:mm) is "
     "reused by storage, the duplicate-slot lookup, and the views, guaranteeing that slot "
     "comparisons match exactly.")

# ======================= 6. USE CASE DIAGRAM =======================
h1("6. Use Case Diagram")
para("Figure 1 presents the use case diagram for MEDIS. It captures the three actors and "
     "the functions each may perform within the system boundary.")
figure("usecase.png", "Figure 1 — MEDIS Use Case Diagram", width=4.0)
para("Actors:", bold=True, space_after=2)
bullet("Admin — privileged staff with access to all four functional areas: managing patients, managing doctors, booking/viewing appointments, and viewing reports (dashboard shows all four tabs).")
bullet("Receptionist — front-desk staff who registers/updates patients and books/views appointments, but does not access reports (Patients + Appointments tabs).")
bullet("Doctor — clinical staff who reviews their appointment list and patient records (Appointments + Patients tabs).")
para("Use cases and rationale:", bold=True, space_after=2)
table(["Use Case", "Description", "Actors"],
      [["Login", "Authenticate; role determines available functions", "All"],
       ["Logout", "End session and return to login screen", "All"],
       ["Add / Update Patient", "Create or modify a patient record", "Admin, Receptionist"],
       ["View Patient Records", "Browse the patient table / history", "Admin, Doctor, Receptionist"],
       ["Add / View Doctors", "Register a doctor and assign a specialisation", "Admin"],
       ["Book Appointment", "Select patient, doctor and date-time; prevents duplicates", "Admin, Receptionist"],
       ["View Appointment List", "View all appointments in a table", "Admin, Doctor, Receptionist"],
       ["Cancel Appointment", "Soft-cancel a selected appointment", "Admin, Receptionist"],
       ["View Reports", "Display totals and doctor schedules", "Admin"]],
      widths=[1.7, 3.1, 1.7])
para("Relationships:", bold=True, space_after=2)
bullet("<<include>> — Add/Update Patient, Book Appointment, and View Reports all include Login, because authentication is a mandatory precondition for every protected action.")
bullet("<<extend>> — Cancel Appointment extends View Appointment List, because cancelling is an optional action on a row already shown in the list (AppointmentPanel.doCancel() acts on the selected row).")

# ======================= 7. CLASS DIAGRAM =======================
h1("7. Class Diagram")
para("Figure 2 shows the class diagram of the implemented system, grouped by architectural "
     "layer (model = blue, controller = green, DAO interface = red, DAO implementation = "
     "pink, infrastructure/util = light blue).")
figure("class.png", "Figure 2 — MEDIS Class Diagram", width=6.5)
para("Required classes (per the assignment) and how they appear:", bold=True, space_after=2)
bullet("User (abstract) — base class holding userId, username, passwordHash, role, with the abstract operation dashboardTitle().")
bullet("Admin, Receptionist, Doctor — concrete subclasses of User.")
bullet("Patient — entity with patientId, name, age, gender, medicalHistory.")
bullet("Appointment — entity holding references to a Patient and a Doctor, plus dateTime and a Status enumeration.")
bullet("SystemController — the central controller aggregating all sub-controllers and the current User.")
para("Relationships and design rationale:", bold=True, space_after=2)
table(["Relationship", "Type", "Meaning"],
      [["User → Admin / Receptionist / Doctor", "Inheritance (generalization)",
        "The three roles specialise abstract User; each overrides dashboardTitle()."],
       ["Appointment → Patient, Appointment → Doctor", "Association (1)",
        "Each appointment references exactly one patient and one doctor (object fields)."],
       ["Appointment ◆— Status", "Composition",
        "The Status enumeration is owned by, and only meaningful within, an Appointment."],
       ["SystemController ◇— sub-controllers", "Aggregation",
        "SystemController holds the five controllers, assembled at startup."],
       ["SystemController → User", "Association", "Holds the current logged-in user (currentUser)."],
       ["*DAOImpl ..▷ *DAO", "Realization", "Each implementation realises its DAO interface."],
       ["Controller ⇢ DAO", "Dependency", "Controllers depend on DAO interfaces; AppointmentController depends on three DAOs."],
       ["*DAOImpl ⇢ DatabaseConnection", "Dependency", "Implementations obtain the shared JDBC connection from the singleton."]],
      widths=[2.3, 1.7, 2.5])
para("A notable design point is that Doctor extends User while also serving as the doctor "
     "entity (it carries doctorId, name, specialization). It provides two constructors — "
     "one for a doctor with a login account and one for a doctor record without one — which "
     "is why UserDAOImpl performs a users LEFT JOIN doctors query and instantiates the "
     "correct subtype based on the role column.")

# ======================= 8. SEQUENCE DIAGRAMS =======================
h1("8. Sequence Diagrams")
para("The four sequence diagrams below correspond exactly to the four flows required by "
     "the assignment (Login, Book Appointment, Add Patient, Generate Report). Each diagram "
     "traces the real method calls in the submitted code, preserving the View → Controller "
     "→ DAO → Database direction and the Result return contract.")

h2("8.1 Login Process")
figure("seq_login.png", "Figure 3 — Sequence Diagram: Login Process", width=6.5)
para("When the user submits credentials, LoginFrame.doLogin() calls "
     "LoginController.authenticate(). The controller first validates that the fields are "
     "not blank (Validator.isNotBlank), then asks UserDAO.findByUsername() to load the "
     "matching record. UserDAOImpl runs a SELECT ... LEFT JOIN doctors query and "
     "instantiates the correct User subclass (Admin, Doctor, or Receptionist) according "
     "to the role column — the system's primary example of polymorphism. The controller "
     "then verifies the password against the stored hash via PasswordHasher.verify(). On "
     "success it returns Result.ok(user); LoginFrame stores the user in SystemController "
     "and opens the DashboardFrame. On failure, DialogHelper.showError is shown and the "
     "user stays on the login screen.")

h2("8.2 Book Appointment")
figure("seq_book.png", "Figure 4 — Sequence Diagram: Book Appointment", width=6.5)
para("AppointmentPanel.doBook() parses the date-time string using the canonical FMT format "
     "and calls AppointmentController.book(). The controller enforces the business rules in "
     "order: the date must be in the future (Validator.isFutureDate), the patient and "
     "doctor must exist (findById), and the slot must be free (existsByDoctorAndTime). Only "
     "then is AppointmentDAO.insert() invoked. Duplicate slots are defended twice — by the "
     "controller's existsByDoctorAndTime check and by a UNIQUE(doctor_id, "
     "appointment_datetime) constraint in the schema, which causes the insert to return "
     "Result.fail(\"Time slot already taken\"). Any rule violation short-circuits with a "
     "Result.fail, displayed via DialogHelper.showError.")

h2("8.3 Add Patient")
figure("seq_addpatient.png", "Figure 5 — Sequence Diagram: Add Patient", width=6.3)
para("PatientPanel.doAdd() parses the age field (rejecting non-numeric input) and calls "
     "PatientController.add(). The controller validates the name, age range (0–150) and "
     "gender, then calls PatientDAO.insert(), which performs an INSERT with "
     "RETURN_GENERATED_KEYS and writes the database-assigned patient_id back onto the "
     "Patient object. A successful Result.ok(patient) triggers an information dialog and a "
     "table refresh; a validation failure returns Result.fail and shows an error.")

h2("8.4 Generate Report")
figure("seq_report.png", "Figure 6 — Sequence Diagram: Generate Report", width=6.3)
para("ReportPanel.refresh() queries ReportController for three figures: totalPatients() "
     "and totalAppointments() (each derived from the corresponding findAll().size()), and "
     "doctorSchedules(), which iterates over every doctor and collects their appointments "
     "via AppointmentDAO.findByDoctor() into a Map<Doctor, List<Appointment>>. The panel "
     "then renders the totals into labels and the schedule into a JTable.")

# ======================= 9. SYSTEM IMPLEMENTATION =======================
h1("9. System Implementation")
para("Technology stack. Java (Swing for UI, JDBC for data access) with an embedded SQLite "
     "database. The SQLite driver and JUnit runner are bundled in lib/, so the project "
     "builds and runs with only a JDK installed — no Maven, Gradle, or network access.")
para("Persistence. db.DatabaseConnection is a singleton holding one shared Connection. On "
     "first launch (when db/medis.db is absent) it executes db/schema.sql to create the "
     "users, patients, doctors, and appointments tables and seed demonstration data. "
     "Foreign keys are enabled (PRAGMA foreign_keys = ON).")
para("Application startup (Main). Main registers the JDBC driver, applies the system "
     "look-and-feel, constructs the DAOs, injects them into the controllers, assembles the "
     "SystemController, and shows the LoginFrame on the Swing event-dispatch thread.")
para("Representative code — duplicate-slot business rule (AppointmentController.book):", bold=True, space_after=2)
code_block(
'public Result book(int patientId, int doctorId, LocalDateTime dt) {\n'
'    if (!Validator.isFutureDate(dt))\n'
'        return Result.fail("Date must be in the future");\n'
'    Optional<Patient> p = patientDAO.findById(patientId);\n'
'    if (p.isEmpty()) return Result.fail("Please select a patient");\n'
'    Optional<Doctor> d = doctorDAO.findById(doctorId);\n'
'    if (d.isEmpty()) return Result.fail("Please select a doctor");\n'
'    if (apptDAO.existsByDoctorAndTime(doctorId, dt))\n'
'        return Result.fail("Time slot already taken");\n'
'    return apptDAO.insert(new Appointment(p.get(), d.get(), dt));\n'
'}')
para("Representative code — polymorphic user construction (UserDAOImpl.mapUser):", bold=True, space_after=2)
code_block(
'switch (role) {\n'
'    case "ADMIN":        return new Admin(id, un, hash);\n'
'    case "RECEPTIONIST": return new Receptionist(id, un, hash);\n'
'    case "DOCTOR":       return new Doctor(id, un, hash, doctorId, name, specialization);\n'
'    default: throw new SQLException("Unknown role: " + role);\n'
'}')
para("Testing. The project includes JUnit tests for Validator and PasswordHasher, and an "
     "end-to-end SmokeTest that drives the controller layer exactly as the GUI does — "
     "covering login, role authorisation, CRUD, future-date and duplicate-slot rejection, "
     "soft-cancel, and report totals.")

# ======================= 10. OOP CONCEPTS APPLIED =======================
h1("10. OOP Concepts Applied")
para("Each of the four required OOP pillars is demonstrated directly in the MEDIS source code.")
h2("10.1 Encapsulation")
para("All entity fields are declared private (or protected in User) and exposed only "
     "through accessor/mutator methods. For example, Patient keeps patientId, name, age, "
     "gender, and medicalHistory private and provides getName(), setAge(), etc. The Result "
     "class goes further: its fields are private final and instances can only be created "
     "through the static factory methods ok() and fail(), making it immutable. The "
     "DatabaseConnection singleton hides its Connection and constructor, exposing only "
     "getInstance() and get().")
h2("10.2 Inheritance")
para("User is an abstract base class providing shared state (userId, username, "
     "passwordHash, role) and behaviour. Admin, Receptionist, and Doctor each extend User, "
     "inheriting that state and supplying role-specific behaviour through their "
     "constructors and overrides. This eliminates duplication across the three roles.")
code_block(
'public abstract class User {\n'
'    protected int userId; protected String username, passwordHash, role;\n'
'    public abstract String dashboardTitle();\n'
'}\n'
'public class Admin extends User {\n'
'    public Admin(int id, String u, String h) { super(id, u, h, "ADMIN"); }\n'
'    @Override public String dashboardTitle() { return "Admin Dashboard"; }\n'
'}')
h2("10.3 Polymorphism")
para("Polymorphism appears in two concrete forms:")
para("1. Subtype polymorphism via the abstract method. dashboardTitle() is declared "
     "abstract in User and overridden by each subclass. DashboardFrame calls "
     "system.getCurrentUser().dashboardTitle() without knowing the concrete type — the "
     "correct title is resolved at runtime.")
para("2. Runtime object substitution in the DAO. UserDAOImpl.findByUsername() returns an "
     "Optional<User> whose actual content may be an Admin, Doctor, or Receptionist, "
     "selected at runtime from the role column. Callers treat all three uniformly through "
     "the User reference.")
para("Method overriding of toString() in Doctor and Patient (used to render combo-box "
     "items) is a further example.")
h2("10.4 Abstraction")
para("Abstraction is realised through interfaces in the DAO layer. PatientDAO, DoctorDAO, "
     "AppointmentDAO, and UserDAO declare what data operations exist (insert, findById, "
     "findAll, update, …) while the …Impl classes encapsulate how they are done with JDBC. "
     "Controllers program against the interface only:")
code_block(
'public class PatientController {\n'
'    private final PatientDAO dao;            // depends on the abstraction\n'
'    public PatientController(PatientDAO dao) { this.dao = dao; }\n'
'}')
para("This decouples business logic from the persistence technology and is the seam that "
     "would allow SQLite to be swapped for another store without changing the controllers. "
     "The abstract User class is a second form of abstraction, modelling the general "
     "concept of a system user independently of any specific role.")

# ======================= 11. SYSTEM FEATURES =======================
h1("11. System Features")
table(["Feature", "Description"],
      [["Role-based access control", "After login, DashboardFrame shows only the tabs permitted for the user's role (Admin: all four; Doctor: Appointments + Patients; Receptionist: Patients + Appointments)."],
       ["Secure authentication", "Passwords are SHA-256 hashed; only hashes are stored and compared."],
       ["Input validation", "Blank-field, age-range (0–150), numeric-age, future-date, and date-format checks reject bad input with clear messages."],
       ["Duplicate-slot prevention", "Enforced both in the controller (existsByDoctorAndTime) and by a database UNIQUE constraint."],
       ["Soft cancellation", "Cancelling an appointment sets its status to CANCELLED; the record is retained for auditability rather than deleted."],
       ["Persistent storage", "All data is stored in SQLite and survives restarts; the database is auto-created and seeded on first run."],
       ["Consistent error handling", "The uniform Result contract drives every success/error dialog through DialogHelper."],
       ["Reporting", "Live totals for patients and appointments plus a per-doctor schedule table."],
       ["Event-driven UI", "All actions are wired through ActionListener lambdas."]],
      widths=[2.0, 4.3])

# ======================= 12. CONCLUSION =======================
h1("12. Conclusion")
para("MEDIS fulfils every functional and technical requirement of the CCP6224 lab "
     "exercise. The five mandated modules are implemented and verified, the GUI is built "
     "entirely in Java Swing with event-driven controls, and persistence is provided by an "
     "embedded SQLite database. The system demonstrates all four object-oriented pillars in "
     "code that is directly traceable to the UML models presented here — abstract User and "
     "DAO interfaces (abstraction), the role hierarchy (inheritance), runtime role "
     "resolution and overridden operations (polymorphism), and private state with "
     "controlled access (encapsulation).")
para("The deliberately clean, layered architecture (View → Controller → DAO → Database) "
     "with a uniform Result contract and a single composition root makes the system "
     "coherent, testable, and extensible. Because every diagram in this report was produced "
     "from the submitted source, the design and implementation remain fully consistent — "
     "satisfying the Design Coherence criterion of the marking rubric.")
para("Possible future enhancements include an in-application user-management screen for "
     "admins, finer-grained appointment statuses (e.g. automatic completion), and "
     "exportable PDF reports.")

# ======================= 13. REFERENCES =======================
h1("13. References")
refs = [
 "Oracle. Java Platform, Standard Edition Documentation — Java Swing (javax.swing). https://docs.oracle.com/en/java/",
 "SQLite Consortium. SQLite Documentation. https://www.sqlite.org/docs.html",
 "Xerial. sqlite-jdbc Driver. https://github.com/xerial/sqlite-jdbc",
 "Object Management Group. Unified Modeling Language (UML) Specification, v2.5.1. https://www.omg.org/spec/UML/",
 "Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). Design Patterns: Elements of Reusable Object-Oriented Software. Addison-Wesley. (Singleton, DAO patterns.)",
 "CCP6224 Object-Oriented Analysis and Design — Lab Exercise 2026 assignment brief and marking rubric.",
]
for i, r in enumerate(refs, 1):
    doc.add_paragraph(f"{i}. {r}", style="List Number" if False else None)

# ======================= 14. APPENDICES =======================
h1("14. Appendices")
h2("Appendix A — Build & Run Instructions")
code_block("# Compile all sources into bin/\n"
           "javac -d bin -cp \"lib/*\" $(find src -name \"*.java\")\n"
           "# Run the application\n"
           "java -cp \"lib/*:bin\" Main\n"
           "# (On Windows, use ; instead of : in the classpath.)")
h2("Appendix B — Default Login Credentials (seed data)")
table(["Role", "Username", "Password"],
      [["Admin", "admin", "admin123"],
       ["Doctor", "doctor1", "pass123"],
       ["Receptionist", "recep1", "pass123"]],
      widths=[2.0, 2.0, 2.0])
h2("Appendix C — Database Schema (tables)")
bullet("users(user_id, username, password, role)")
bullet("patients(patient_id, name, age, gender, medical_history, created_at)")
bullet("doctors(doctor_id, name, specialization, user_id)")
bullet("appointments(appointment_id, patient_id, doctor_id, appointment_datetime, status, UNIQUE(doctor_id, appointment_datetime))")
h2("Appendix D — Automated Tests")
bullet("JUnit unit tests: util.ValidatorTest, util.PasswordHasherTest.")
bullet("End-to-end controller smoke test: SmokeTest (login, RBAC, CRUD, duplicate-slot & future-date rejection, soft-cancel, report totals).")

# ---------- footer with page numbers + update-fields-on-open ----------
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.add_run("Page ")
add_field(fp, "PAGE")
fp.add_run(" of ")
add_field(fp, "NUMPAGES")

settings = doc.settings.element
upd = OxmlElement("w:updateFields"); upd.set(qn("w:val"), "true")
settings.append(upd)

doc.save(OUT)
print("Saved:", OUT)
